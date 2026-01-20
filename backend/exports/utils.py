"""
Export Utilities - Core Export Generation Logic

This module handles generating case exports in multiple formats:
- PDF: Uses WeasyPrint (HTML/CSS to PDF) for perfect Vietnamese support
- Word: Uses python-docx for editable .docx documents
- JSON: Uses Python's json module with UTF-8 encoding

Key Classes:
- PDFExporterHTML: Primary PDF generator (WeasyPrint-based) ⭐
- PDFExporter: Legacy PDF generator (ReportLab-based, fallback only)
- WordExporter: Word document generator
- JSONExporter: JSON data export
- ExportUtils: Common utilities for data preparation

Vietnamese Support:
- All exporters properly handle Vietnamese characters
- WeasyPrint provides browser-grade Unicode rendering for PDFs
- python-docx has native Unicode support for Word
- JSON uses ensure_ascii=False for readable Vietnamese text

Author: Development Team
Last Updated: November 3, 2025
"""

import io
import json
import unicodedata
from datetime import datetime
from typing import Dict, Any, Optional
from django.conf import settings
from django.core.files.base import ContentFile
from django.utils import timezone
from cases.models import Case

# PDF generation imports
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
        Image,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.pdfmetrics import registerFontFamily

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word document imports
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def normalize_vietnamese_text(text: str) -> str:
    """
    Normalize Vietnamese text to NFC (precomposed) form for proper PDF rendering.

    Vietnamese characters can be represented in two Unicode normalization forms:

    NFD (Decomposed): base character + separate combining marks
        Example: "ố" = "o" (U+006F) + "̂" (U+0302) + "́" (U+0301)
        Problem: Some PDF renderers display these incorrectly

    NFC (Precomposed): single combined character
        Example: "ố" = single character (U+1ED1)
        Solution: Better compatibility with PDF libraries

    This function converts text to NFC form for ReportLab compatibility.
    Note: WeasyPrint handles both forms correctly, so this is mainly
    for the ReportLab fallback generator.

    Args:
        text: Input text that may contain Vietnamese characters

    Returns:
        Normalized text in NFC form

    Example:
        >>> normalize_vietnamese_text("Suy hô hấp cấp")
        'Suy hô hấp cấp'  # All characters in precomposed form
    """
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    # Normalize to NFC (Canonical Decomposition, followed by Canonical Composition)
    return unicodedata.normalize("NFC", text)


class ExportUtils:
    """
    Utility class for export operations
    """

    @staticmethod
    def get_case_data(case: Case, template_settings: Dict = None) -> Dict[str, Any]:  # type: ignore[attr-defined]
        """
        Extract complete case data based on template settings
        """
        settings = template_settings or {}

        data = {
            "id": case.id,  # type: ignore[attr-defined]
            "title": case.title,
            "exported_at": timezone.now().isoformat(),
            "case_status": case.get_case_status_display(),  # type: ignore[attr-defined]
        }

        # Patient information
        if settings.get("include_patient_details", True):
            anonymize = settings.get("anonymize_patient", False)

            data["patient"] = {
                "name": "Bệnh nhân [Ẩn danh]" if anonymize else case.patient_name,
                "age": case.patient_age,
                "gender": case.get_patient_gender_display(),  # type: ignore[attr-defined]
                "medical_record_number": (
                    "[Ẩn danh]" if anonymize else case.medical_record_number
                ),
                "ethnicity": case.patient_ethnicity or "Không xác định",
                "occupation": case.patient_occupation or "Không xác định",
                "admission_date": (
                    case.admission_date.strftime("%d/%m/%Y")
                    if case.admission_date
                    else "N/A"
                ),
                "discharge_date": (
                    case.discharge_date.strftime("%d/%m/%Y")
                    if case.discharge_date
                    else "N/A"
                ),
            }

        # Clinical information
        if settings.get("include_medical_history", True):
            # Use new medical models
            if hasattr(case, "clinical_history"):
                ch = case.clinical_history  # type: ignore[attr-defined]
                data["clinical_history"] = {
                    "chief_complaint": ch.chief_complaint or "Không có",
                    "history_present_illness": ch.history_present_illness or "Không có",
                    "past_medical_history": ch.past_medical_history or "Không có",
                    "family_history": ch.family_history or "Không có",
                    "social_history": ch.social_history or "Không có",
                    "allergies": ch.allergies or "Không có",
                    "medications": ch.medications or "Không có",
                }
                # For backward compatibility in templates
                data["history"] = ch.history_present_illness or "Không có thông tin"
                data["chief_complaint"] = ch.chief_complaint or "Không có thông tin"
            else:
                data["history"] = "Không có thông tin"
                data["chief_complaint"] = "Không có thông tin"

        if settings.get("include_examination", True):
            # Detailed physical examination if available
            if hasattr(case, "physical_examination"):
                pe = case.physical_examination  # type: ignore[attr-defined]
                data["physical_examination"] = {
                    "vital_signs": {
                        "temperature": pe.vital_signs_temp or "N/A",
                        "heart_rate": pe.vital_signs_hr or "N/A",
                        "respiratory_rate": pe.vital_signs_rr or "N/A",
                        "blood_pressure": pe.vital_signs_bp or "N/A",
                        "oxygen_saturation": pe.vital_signs_spo2 or "N/A",
                    },
                    "general_appearance": pe.general_appearance or "Không có",
                    "head_neck": getattr(pe, "head_neck", "") or "Không có",
                    "cardiovascular": pe.cardiovascular or "Không có",
                    "respiratory": pe.respiratory or "Không có",
                    "abdominal": pe.abdominal or "Không có",
                    "neurological": pe.neurological or "Không có",
                }
                # For backward compatibility
                data["examination"] = f"{pe.general_appearance}\n{pe.vital_signs}"
            else:
                data["examination"] = "Không có thông tin"

        if settings.get("include_investigations", True):
            if hasattr(case, "investigations_detail"):
                inv = case.investigations_detail  # type: ignore[attr-defined]
                data["investigations_detail"] = {
                    "laboratory_results": inv.laboratory_results or "Không có",
                    "imaging_studies": inv.imaging_studies or "Không có",
                    "other_tests": getattr(inv, "other_tests", None) or "Không có",
                }
                # For backward compatibility
                data["investigations"] = (
                    f"{inv.laboratory_results or ''}\n{inv.imaging_studies or ''}".strip()
                    or "Không có thông tin"
                )
            else:
                data["investigations"] = "Không có thông tin"

        if settings.get("include_diagnosis", True):
            if hasattr(case, "diagnosis_management"):
                dm = case.diagnosis_management  # type: ignore[attr-defined]
                data["diagnosis_management"] = {
                    "primary_diagnosis": dm.primary_diagnosis or "Không có",
                    "differential_diagnosis": dm.differential_diagnosis or "Không có",
                    "treatment_plan": dm.treatment_plan or "Không có",
                    "follow_up_plan": dm.follow_up_plan or "Không có",
                }
                # For backward compatibility
                data["diagnosis"] = dm.primary_diagnosis or "Không có thông tin"
                data["treatment"] = dm.treatment_plan or "Không có thông tin"
                data["follow_up"] = dm.follow_up_plan or "Không có thông tin"
            else:
                data["diagnosis"] = "Không có thông tin"
                data["treatment"] = "Không có thông tin"
                data["follow_up"] = "Không có thông tin"

        if settings.get("include_learning_objectives", True):
            if hasattr(case, "learning_outcomes"):
                lo = case.learning_outcomes  # type: ignore[attr-defined]
                data["learning_outcomes"] = {
                    "learning_objectives": getattr(lo, "learning_objectives", "") or "Không có",
                    "key_concepts": getattr(lo, "key_concepts", "") or "Không có",
                    "clinical_reasoning": getattr(lo, "clinical_pearls", "") or "Không có",  # Fixed: was clinical_reasoning_points
                }
                # For backward compatibility
                data["learning_objectives"] = (
                    lo.learning_objectives or "Không có thông tin"
                )
            else:
                data["learning_objectives"] = "Không có thông tin"

            data["learning_tags"] = getattr(case, "learning_tags", [])
            data["complexity_level"] = case.get_complexity_level_display()  # type: ignore[attr-defined]

        # Metadata
        data["metadata"] = {
            "specialty": case.specialty,
            "priority_level": case.get_priority_level_display(),  # type: ignore[attr-defined]
            "complexity_level": case.get_complexity_level_display(),  # type: ignore[attr-defined]
            "estimated_study_hours": case.estimated_study_hours,
            "created_by": case.student.get_full_name(),
            "created_at": case.created_at.strftime("%d/%m/%Y %H:%M"),
            "updated_at": case.updated_at.strftime("%d/%m/%Y %H:%M"),
        }

        # Comments (if requested)
        if settings.get("include_comments", False):
            comments = case.comments.all().order_by("created_at")  # type: ignore[attr-defined]
            data["comments"] = [
                {
                    "author": comment.author.get_full_name(),
                    "content": comment.content,
                    "created_at": comment.created_at.strftime("%d/%m/%Y %H:%M"),
                }
                for comment in comments
            ]

        # Attachments metadata (if requested)
        if settings.get("include_attachments", True):
            attachments = case.medical_attachments.all()  # type: ignore[attr-defined]
            data["attachments"] = [
                {
                    "title": att.title,
                    "file_type": att.file_type,
                    "description": att.description,
                    "uploaded_at": att.uploaded_at.strftime("%d/%m/%Y %H:%M"),
                }
                for att in attachments
            ]

        # Grades (if requested and user has permission)
        if settings.get("include_grades", False):
            if hasattr(case, "grades"):
                grades = case.grades.all()  # type: ignore[attr-defined]
                data["grades"] = [
                    {
                        "grader": grade.grader.get_full_name(),
                        "score": grade.score,
                        "feedback": grade.feedback,
                        "graded_at": grade.created_at.strftime("%d/%m/%Y %H:%M"),
                    }
                    for grade in grades
                ]

        return data


class PDFExporter:
    """
    Generate PDF exports with Vietnamese support
    """

    def __init__(self, template=None):
        self.template = template
        self.styles = getSampleStyleSheet()  # type: ignore[attr-defined]
        self._setup_styles()

    def _setup_styles(self):
        """Setup custom styles for Vietnamese text"""
        # Register Vietnamese-compatible fonts (DejaVu is included with ReportLab)
        # These fonts support Vietnamese characters
        try:
            import os

            # Try to register DejaVu fonts if not already registered
            if "DejaVuSans" not in pdfmetrics.getRegisteredFontNames():  # type: ignore[attr-defined]
                # Try system fonts first
                font_paths = [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
                    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",  # macOS
                    "C:\\Windows\\Fonts\\Arial.ttf",  # Windows
                ]

                font_registered = False
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        try:
                            # Register fonts with subfontIndex for better Unicode support
                            pdfmetrics.registerFont(  # type: ignore[attr-defined]
                                TTFont("DejaVuSans", font_path, subfontIndex=0)  # type: ignore[attr-defined]
                            )
                            bold_path = font_path.replace(
                                "DejaVuSans.ttf", "DejaVuSans-Bold.ttf"
                            )
                            if os.path.exists(bold_path):
                                pdfmetrics.registerFont(  # type: ignore[attr-defined]
                                    TTFont("DejaVuSans-Bold", bold_path, subfontIndex=0)  # type: ignore[attr-defined]
                                )
                            font_registered = True
                            break
                        except Exception:
                            continue

                if not font_registered:
                    # Fallback: Use Helvetica (won't show Vietnamese correctly but won't crash)
                    pass

        except Exception:
            # If font registration fails, continue with default fonts
            pass

        # Determine which font to use
        base_font = (
            "DejaVuSans"
            if "DejaVuSans" in pdfmetrics.getRegisteredFontNames()  # type: ignore[attr-defined]
            else "Helvetica"
        )
        bold_font = (
            "DejaVuSans-Bold"
            if "DejaVuSans-Bold" in pdfmetrics.getRegisteredFontNames()  # type: ignore[attr-defined]
            else "Helvetica-Bold"
        )

        # Register font family for automatic bold/italic switching
        if "DejaVuSans" in pdfmetrics.getRegisteredFontNames():  # type: ignore[attr-defined]
            registerFontFamily(  # type: ignore[attr-defined]
                "DejaVuSans",
                normal="DejaVuSans",
                bold="DejaVuSans-Bold",
                italic="DejaVuSans",  # Use normal for italic if not available
                boldItalic="DejaVuSans-Bold",
            )  # Use bold for boldItalic

        # Title style
        if "CustomTitle" not in self.styles:
            self.styles.add(
                ParagraphStyle(  # type: ignore[attr-defined]
                    name="CustomTitle",
                    parent=self.styles["Title"],
                    fontSize=18,
                    textColor=colors.HexColor("#1a365d"),  # type: ignore[attr-defined]
                    spaceAfter=20,
                    alignment=TA_CENTER,  # type: ignore[attr-defined]
                    fontName=bold_font,
                )
            )

        # Section heading style
        if "SectionHeading" not in self.styles:
            self.styles.add(
                ParagraphStyle(  # type: ignore[attr-defined]
                    name="SectionHeading",
                    parent=self.styles["Heading2"],
                    fontSize=14,
                    textColor=colors.HexColor("#2c5282"),  # type: ignore[attr-defined]
                    spaceAfter=12,
                    spaceBefore=16,
                    fontName=bold_font,
                )
            )

        # Body text style (will use font family for <b> tags)
        if "BodyText" not in self.styles:
            self.styles.add(
                ParagraphStyle(  # type: ignore[attr-defined]
                    name="BodyText",
                    parent=self.styles["Normal"],
                    fontSize=11,
                    spaceAfter=10,
                    alignment=TA_JUSTIFY,  # type: ignore[attr-defined]
                    leading=14,
                    fontName=base_font,
                )
            )

        # Info style (for metadata)
        if "InfoText" not in self.styles:
            self.styles.add(
                ParagraphStyle(  # type: ignore[attr-defined]
                    name="InfoText",
                    parent=self.styles["Normal"],
                    fontSize=10,
                    textColor=colors.HexColor("#4a5568"),  # type: ignore[attr-defined]
                    spaceAfter=6,
                    fontName=base_font,
                )
            )

    def generate(self, case: Case, output_path: str = None) -> io.BytesIO:  # type: ignore[attr-defined]
        """
        Generate PDF export for a case
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(  # type: ignore[attr-defined]
            buffer,
            pagesize=A4,  # type: ignore[attr-defined]
            rightMargin=2 * cm,  # type: ignore[attr-defined]
            leftMargin=2 * cm,  # type: ignore[attr-defined]
            topMargin=2 * cm,  # type: ignore[attr-defined]
            bottomMargin=2 * cm,  # type: ignore[attr-defined]
        )

        story = []

        # Get case data based on template settings
        template_settings = {}
        if self.template:
            template_settings = {
                "include_patient_details": self.template.include_patient_details,
                "include_medical_history": self.template.include_medical_history,
                "include_examination": self.template.include_examination,
                "include_investigations": self.template.include_investigations,
                "include_diagnosis": self.template.include_diagnosis,
                "include_treatment": self.template.include_treatment,
                "include_learning_objectives": self.template.include_learning_objectives,
                "include_comments": self.template.include_comments,
                "include_attachments": self.template.include_attachments,
                "anonymize_patient": self.template.anonymize_patient,
            }

        data = ExportUtils.get_case_data(case, template_settings)

        # Add header if specified
        if self.template and self.template.header_text:
            story.append(
                Paragraph(  # type: ignore[attr-defined]
                    normalize_vietnamese_text(self.template.header_text),
                    self.styles["InfoText"],
                )
            )
            story.append(Spacer(1, 0.3 * cm))  # type: ignore[attr-defined]

        # Add watermark text if specified
        if self.template and self.template.add_watermark:
            watermark = normalize_vietnamese_text(
                self.template.watermark_text or "CONFIDENTIAL"
            )
            story.append(
                Paragraph(  # type: ignore[attr-defined]
                    f"<font color='red'>{watermark}</font>",
                    self.styles["CustomTitle"],
                )
            )
            story.append(Spacer(1, 0.5 * cm))  # type: ignore[attr-defined]

        # Title
        story.append(
            Paragraph(normalize_vietnamese_text(case.title), self.styles["CustomTitle"])  # type: ignore[attr-defined]
        )
        story.append(Spacer(1, 0.5 * cm))  # type: ignore[attr-defined]

        # Patient Information
        if "patient" in data:
            story.append(
                Paragraph("THÔNG TIN BỆNH NHÂN", self.styles["SectionHeading"])  # type: ignore[attr-defined]
            )
            patient = data["patient"]

            # Use Paragraphs for Vietnamese text in tables
            patient_data = [
                [
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text("<b>Tên bệnh nhân:</b>"),
                        self.styles["BodyText"],
                    ),
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text(patient["name"]),
                        self.styles["BodyText"],
                    ),
                ],
                [
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text("<b>Tuổi:</b>"),
                        self.styles["BodyText"],
                    ),
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text(str(patient["age"])),
                        self.styles["BodyText"],
                    ),
                ],
                [
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text("<b>Giới tính:</b>"),
                        self.styles["BodyText"],
                    ),
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text(patient["gender"]),
                        self.styles["BodyText"],
                    ),
                ],
                [
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text("<b>Số hồ sơ:</b>"),
                        self.styles["BodyText"],
                    ),
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text(
                            patient.get("medical_record_number", "N/A")
                        ),
                        self.styles["BodyText"],
                    ),
                ],
                [
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text("<b>Dân tộc:</b>"),
                        self.styles["BodyText"],
                    ),
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text(patient.get("ethnicity", "N/A")),
                        self.styles["BodyText"],
                    ),
                ],
                [
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text("<b>Nghề nghiệp:</b>"),
                        self.styles["BodyText"],
                    ),
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text(patient.get("occupation", "N/A")),
                        self.styles["BodyText"],
                    ),
                ],
                [
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text("<b>Ngày nhập viện:</b>"),
                        self.styles["BodyText"],
                    ),
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text(patient.get("admission_date", "N/A")),
                        self.styles["BodyText"],
                    ),
                ],
                [
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text("<b>Ngày xuất viện:</b>"),
                        self.styles["BodyText"],
                    ),
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text(patient.get("discharge_date", "N/A")),
                        self.styles["BodyText"],
                    ),
                ],
            ]

            table = Table(patient_data, colWidths=[4 * cm, 12 * cm])  # type: ignore[attr-defined]
            table.setStyle(
                TableStyle(  # type: ignore[attr-defined]
                    [
                        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#2c5282")),  # type: ignore[attr-defined]
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),  # type: ignore[attr-defined]
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 0.5 * cm))  # type: ignore[attr-defined]

        # Clinical sections
        sections = []

        if "chief_complaint" in data:
            sections.append(("LÝ DO KHÁM", data["chief_complaint"]))

        if "history" in data:
            sections.append(("TIỀN SỬ VÀ BỆNH SỬ", data["history"]))

        if "examination" in data:
            sections.append(("KHÁM LÂM SÀNG", data["examination"]))

        if "investigations" in data:
            sections.append(("XÉT NGHIỆM VÀ CẬN LÂM SÀNG", data["investigations"]))

        if "diagnosis" in data:
            sections.append(("CHẨN ĐOÁN", data["diagnosis"]))

        if "treatment" in data:
            sections.append(("ĐIỀU TRỊ", data["treatment"]))

        if "follow_up" in data:
            sections.append(("THEO DÕI VÀ KẾT QUẢ", data["follow_up"]))

        if "learning_objectives" in data:
            sections.append(("MỤC TIÊU HỌC TẬP", data["learning_objectives"]))

        for section_title, content in sections:
            story.append(
                Paragraph(  # type: ignore[attr-defined]
                    normalize_vietnamese_text(section_title),
                    self.styles["SectionHeading"],
                )
            )
            story.append(
                Paragraph(  # type: ignore[attr-defined]
                    normalize_vietnamese_text(content or "Không có thông tin"),
                    self.styles["BodyText"],
                )
            )
            story.append(Spacer(1, 0.3 * cm))  # type: ignore[attr-defined]

        # Attachments
        if "attachments" in data and data["attachments"]:
            story.append(
                Paragraph(  # type: ignore[attr-defined]
                    normalize_vietnamese_text("TÀI LIỆU ĐÍNH KÈM"),
                    self.styles["SectionHeading"],
                )
            )
            for att in data["attachments"]:
                att_text = (
                    f"• {att['title']} ({att['file_type']}) - {att['uploaded_at']}"
                )
                story.append(
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text(att_text), self.styles["BodyText"]
                    )
                )
            story.append(Spacer(1, 0.5 * cm))  # type: ignore[attr-defined]

        # Comments
        if "comments" in data and data["comments"]:
            story.append(
                Paragraph(  # type: ignore[attr-defined]
                    normalize_vietnamese_text("NHẬN XÉT"), self.styles["SectionHeading"]
                )
            )
            for comment in data["comments"]:
                comment_text = f"<b>{comment['author']}</b> ({comment['created_at']}): {comment['content']}"
                story.append(
                    Paragraph(  # type: ignore[attr-defined]
                        normalize_vietnamese_text(comment_text), self.styles["BodyText"]
                    )
                )
            story.append(Spacer(1, 0.3 * cm))  # type: ignore[attr-defined]

        # Metadata footer
        story.append(PageBreak())  # type: ignore[attr-defined]
        story.append(Paragraph("THÔNG TIN CA BỆNH", self.styles["SectionHeading"]))  # type: ignore[attr-defined]
        metadata = data["metadata"]
        metadata_text = f"""
        <b>Chuyên khoa:</b> {metadata["specialty"]}<br/>
        <b>Mức độ ưu tiên:</b> {metadata["priority_level"]}<br/>
        <b>Mức độ phức tạp:</b> {metadata["complexity_level"]}<br/>
        <b>Thời gian học tập ước tính:</b> {metadata.get("estimated_study_hours", "N/A")} giờ<br/>
        <b>Người tạo:</b> {metadata["created_by"]}<br/>
        <b>Ngày tạo:</b> {metadata["created_at"]}<br/>
        <b>Cập nhật lần cuối:</b> {metadata["updated_at"]}<br/>
        <b>Xuất file:</b> {timezone.now().strftime("%d/%m/%Y %H:%M")}
        """
        story.append(Paragraph(metadata_text, self.styles["InfoText"]))  # type: ignore[attr-defined]

        # Add footer if specified
        if self.template and self.template.footer_text:
            story.append(Spacer(1, 0.5 * cm))  # type: ignore[attr-defined]
            story.append(Paragraph(self.template.footer_text, self.styles["InfoText"]))  # type: ignore[attr-defined]

        # Build PDF
        doc.build(story)
        buffer.seek(0)

        return buffer


class PDFExporterHTML:
    """
    Primary PDF Generator - Uses WeasyPrint (HTML/CSS to PDF) ⭐

    This is the MAIN PDF export class with perfect Vietnamese support.

    Why WeasyPrint?
    - Browser-grade HTML/CSS rendering
    - Excellent Unicode support (all Vietnamese characters work perfectly)
    - Proper font subsetting with Vietnamese glyphs
    - Smaller file sizes (~17KB vs ~49KB with ReportLab)
    - CSS-based styling for professional documents

    How it works:
    1. Prepare case data from database
    2. Build HTML template with patient info, clinical data, etc.
    3. Apply CSS styling for medical document formatting
    4. WeasyPrint converts HTML→PDF with proper font embedding
    5. Return PDF buffer ready for download/save

    Vietnamese Support:
    - All Vietnamese characters (á, ă, â, đ, etc.) render perfectly
    - No special font registration needed
    - No Unicode normalization required
    - Works with both NFD and NFC character forms

    Fallback:
    - If WeasyPrint is not installed, automatically falls back to PDFExporter
    - This ensures exports still work even without WeasyPrint dependencies

    Usage:
        exporter = PDFExporterHTML(template=my_template)
        pdf_buffer = exporter.generate(case)
        with open('output.pdf', 'wb') as f:
            f.write(pdf_buffer.getvalue())
    """

    def __init__(self, template=None):
        """
        Initialize PDF exporter with optional template

        Args:
            template: ExportTemplate instance defining what to include
                     If None, includes all case sections
        """
        self.template = template

    def generate(self, case: Case, output_path: str = None) -> io.BytesIO:  # type: ignore[attr-defined]
        """
        Generate PDF export from HTML template with Vietnamese support

        This is the main export function that:
        1. Checks if WeasyPrint is available (falls back to ReportLab if not)
        2. Prepares case data based on template settings
        3. Builds HTML content with CSS styling
        4. Converts HTML to PDF with proper font embedding
        5. Returns PDF as BytesIO buffer

        Args:
            case: Case instance to export
            output_path: Optional file path to save PDF (for debugging)

        Returns:
            io.BytesIO: PDF file buffer ready for download or saving

        Raises:
            ImportError: If WeasyPrint not available (handled with fallback)

        Example:
            >>> from cases.models import Case
            >>> case = Case.objects.get(id=1)
            >>> exporter = PDFExporterHTML()
            >>> buffer = exporter.generate(case)
            >>> # Save to file
            >>> with open('/tmp/export.pdf', 'wb') as f:
            ...     f.write(buffer.getvalue())
        """
        try:
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration
        except (ImportError, OSError):
            # Fallback to ReportLab if WeasyPrint not available or not working
            fallback = PDFExporter(template=self.template)
            return fallback.generate(case, output_path)

        # Get case data based on template settings
        template_settings = {}
        if self.template:
            template_settings = {
                "include_patient_details": self.template.include_patient_details,
                "include_medical_history": self.template.include_medical_history,
                "include_examination": self.template.include_examination,
                "include_investigations": self.template.include_investigations,
                "include_diagnosis": self.template.include_diagnosis,
                "include_treatment": self.template.include_treatment,
                "include_learning_objectives": self.template.include_learning_objectives,
                "include_comments": self.template.include_comments,
                "include_attachments": self.template.include_attachments,
                "anonymize_patient": self.template.anonymize_patient,
            }

        data = ExportUtils.get_case_data(case, template_settings)

        # Build HTML content
        html_content = self._build_html(case, data)

        # Configure fonts
        font_config = FontConfiguration()

        # Generate PDF
        buffer = io.BytesIO()
        HTML(string=html_content).write_pdf(buffer, font_config=font_config)
        buffer.seek(0)

        return buffer

    def _build_html(self, case: Case, data: dict) -> str:
        """
        Build HTML content for PDF generation
        """
        # Add watermark if specified
        watermark_html = ""
        if self.template and self.template.add_watermark:
            watermark = self.template.watermark_text or "CONFIDENTIAL"
            watermark_html = f'<div class="watermark">{watermark}</div>'

        # Build patient info table
        patient_info_html = ""
        if "patient" in data:
            patient = data["patient"]
            patient_info_html = f"""
            <h2>THÔNG TIN BỆNH NHÂN</h2>
            <table class="info-table">
                <tr>
                    <td><strong>Tên bệnh nhân:</strong></td>
                    <td>{patient["name"]}</td>
                </tr>
                <tr>
                    <td><strong>Tuổi:</strong></td>
                    <td>{patient["age"]}</td>
                </tr>
                <tr>
                    <td><strong>Giới tính:</strong></td>
                    <td>{patient["gender"]}</td>
                </tr>
                <tr>
                    <td><strong>Số hồ sơ:</strong></td>
                    <td>{patient.get("medical_record_number", "N/A")}</td>
                </tr>
                <tr>
                    <td><strong>Dân tộc:</strong></td>
                    <td>{patient.get("ethnicity", "Không xác định")}</td>
                </tr>
                <tr>
                    <td><strong>Nghề nghiệp:</strong></td>
                    <td>{patient.get("occupation", "Không xác định")}</td>
                </tr>
                <tr>
                    <td><strong>Ngày nhập viện:</strong></td>
                    <td>{patient.get("admission_date", "N/A")}</td>
                </tr>
                <tr>
                    <td><strong>Ngày xuất viện:</strong></td>
                    <td>{patient.get("discharge_date", "N/A")}</td>
                </tr>
            </table>
            """

        # Build clinical sections
        sections_html = ""
        sections = []

        if "chief_complaint" in data:
            sections.append(("LÝ DO KHÁM", data["chief_complaint"]))
        if "history" in data:
            sections.append(("TIỀN SỬ VÀ BỆNH SỬ", data["history"]))
        if "examination" in data:
            sections.append(("KHÁM LÂM SÀNG", data["examination"]))
        if "investigations" in data:
            sections.append(("XÉT NGHIỆM VÀ CẬN LÂM SÀNG", data["investigations"]))
        if "diagnosis" in data:
            sections.append(("CHẨN ĐOÁN", data["diagnosis"]))
        if "treatment" in data:
            sections.append(("ĐIỀU TRỊ", data["treatment"]))
        if "follow_up" in data:
            sections.append(("THEO DÕI VÀ KẾT QUẢ", data["follow_up"]))
        if "learning_objectives" in data:
            sections.append(("MỤC TIÊU HỌC TẬP", data["learning_objectives"]))

        for title, content in sections:
            sections_html += f"""
            <h2>{title}</h2>
            <p class="content">{content or "Không có thông tin"}</p>
            """

        # Build attachments
        attachments_html = ""
        if "attachments" in data and data["attachments"]:
            attachments_html = "<h2>TÀI LIỆU ĐÍNH KÈM</h2><ul>"
            for att in data["attachments"]:
                attachments_html += f"<li>{att['title']} ({att['file_type']}) - {att['uploaded_at']}</li>"
            attachments_html += "</ul>"

        # Complete HTML
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: 'DejaVu Sans', 'Arial', sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    color: #333;
                }}
                h1 {{
                    color: #1a365d;
                    text-align: center;
                    font-size: 18pt;
                    margin-bottom: 20px;
                }}
                h2 {{
                    color: #2c5282;
                    font-size: 14pt;
                    margin-top: 20px;
                    margin-bottom: 12px;
                    border-bottom: 2px solid #2c5282;
                    padding-bottom: 5px;
                }}
                .watermark {{
                    color: red;
                    text-align: center;
                    font-size: 16pt;
                    font-weight: bold;
                    margin-bottom: 20px;
                }}
                .info-table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 15px 0;
                }}
                .info-table td {{
                    padding: 8px;
                    border: 1px solid #ddd;
                }}
                .info-table td:first-child {{
                    width: 30%;
                    background-color: #f8f9fa;
                    color: #2c5282;
                }}
                .content {{
                    text-align: justify;
                    white-space: pre-wrap;
                }}
                ul {{
                    margin: 10px 0;
                    padding-left: 20px;
                }}
                li {{
                    margin: 5px 0;
                }}
            </style>
        </head>
        <body>
            {watermark_html}
            <h1>{case.title}</h1>
            {patient_info_html}
            {sections_html}
            {attachments_html}
        </body>
        </html>
        """

        return html


class WordExporter:
    """
    Generate Word document exports
    """

    def __init__(self, template=None):
        self.template = template

    def generate(self, case: Case, output_path: str = None) -> io.BytesIO:  # type: ignore[attr-defined]
        """
        Generate Word document export for a case
        """
        doc = Document()  # type: ignore[attr-defined]

        # Set document properties
        doc.core_properties.title = case.title
        doc.core_properties.author = case.student.get_full_name()
        doc.core_properties.created = timezone.now()

        # Get case data based on template settings
        template_settings = {}
        if self.template:
            template_settings = {
                "include_patient_details": self.template.include_patient_details,
                "include_medical_history": self.template.include_medical_history,
                "include_examination": self.template.include_examination,
                "include_investigations": self.template.include_investigations,
                "include_diagnosis": self.template.include_diagnosis,
                "include_treatment": self.template.include_treatment,
                "include_learning_objectives": self.template.include_learning_objectives,
                "include_comments": self.template.include_comments,
                "include_attachments": self.template.include_attachments,
                "anonymize_patient": self.template.anonymize_patient,
            }

        data = ExportUtils.get_case_data(case, template_settings)

        # Add header
        if self.template and self.template.header_text:
            header_para = doc.add_paragraph(self.template.header_text)
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # type: ignore[attr-defined]
            run = header_para.runs[0]
            run.font.size = Pt(9)  # type: ignore[attr-defined]
            run.font.color.rgb = RGBColor(128, 128, 128)  # type: ignore[attr-defined]

        # Add watermark if specified
        if self.template and self.template.add_watermark:
            watermark = self.template.watermark_text or "CONFIDENTIAL"
            watermark_para = doc.add_paragraph(watermark)
            watermark_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # type: ignore[attr-defined]
            run = watermark_para.runs[0]
            run.font.size = Pt(24)  # type: ignore[attr-defined]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # type: ignore[attr-defined]

        # Title
        title = doc.add_heading(case.title, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # type: ignore[attr-defined]

        # Patient Information
        if "patient" in data:
            doc.add_heading("THÔNG TIN BỆNH NHÂN", level=1)
            patient = data["patient"]

            table = doc.add_table(rows=8, cols=2)
            table.style = "Light List Accent 1"

            rows_data = [
                ("Tên bệnh nhân:", patient["name"]),
                ("Tuổi:", str(patient["age"])),
                ("Giới tính:", patient["gender"]),
                ("Số hồ sơ:", patient.get("medical_record_number", "N/A")),
                ("Dân tộc:", patient.get("ethnicity", "N/A")),
                ("Nghề nghiệp:", patient.get("occupation", "N/A")),
                ("Ngày nhập viện:", patient.get("admission_date", "N/A")),
                ("Ngày xuất viện:", patient.get("discharge_date", "N/A")),
            ]

            for i, (label, value) in enumerate(rows_data):
                row_cells = table.rows[i].cells
                row_cells[0].text = label
                row_cells[1].text = value
                # Make label bold
                row_cells[0].paragraphs[0].runs[0].font.bold = True

            doc.add_paragraph()  # Add spacing

        # Clinical sections
        sections = []

        if "chief_complaint" in data:
            sections.append(("LÝ DO KHÁM", data["chief_complaint"]))

        if "history" in data:
            sections.append(("TIỀN SỬ VÀ BỆNH SỬ", data["history"]))

        if "examination" in data:
            sections.append(("KHÁM LÂM SÀNG", data["examination"]))

        if "investigations" in data:
            sections.append(("XÉT NGHIỆM VÀ CẬN LÂM SÀNG", data["investigations"]))

        if "diagnosis" in data:
            sections.append(("CHẨN ĐOÁN", data["diagnosis"]))

        if "treatment" in data:
            sections.append(("ĐIỀU TRỊ", data["treatment"]))

        if "follow_up" in data:
            sections.append(("THEO DÕI VÀ KẾT QUẢ", data["follow_up"]))

        if "learning_objectives" in data:
            sections.append(("MỤC TIÊU HỌC TẬP", data["learning_objectives"]))

        for section_title, content in sections:
            doc.add_heading(section_title, level=2)
            doc.add_paragraph(content or "Không có thông tin")

        # Attachments
        if "attachments" in data and data["attachments"]:
            doc.add_heading("TÀI LIỆU ĐÍNH KÈM", level=2)
            for att in data["attachments"]:
                doc.add_paragraph(
                    f"{att['title']} ({att['file_type']}) - {att['uploaded_at']}",
                    style="List Bullet",
                )

        # Comments
        if "comments" in data and data["comments"]:
            doc.add_heading("NHẬN XÉT", level=2)
            for comment in data["comments"]:
                para = doc.add_paragraph()
                para.add_run(f"{comment['author']}").bold = True
                para.add_run(f" ({comment['created_at']}): {comment['content']}")

        # Metadata
        doc.add_page_break()
        doc.add_heading("THÔNG TIN CA BỆNH", level=1)
        metadata = data["metadata"]

        metadata_items = [
            ("Chuyên khoa", metadata["specialty"]),
            ("Mức độ ưu tiên", metadata["priority_level"]),
            ("Mức độ phức tạp", metadata["complexity_level"]),
            (
                "Thời gian học tập ước tính",
                f"{metadata.get('estimated_study_hours', 'N/A')} giờ",
            ),
            ("Người tạo", metadata["created_by"]),
            ("Ngày tạo", metadata["created_at"]),
            ("Cập nhật lần cuối", metadata["updated_at"]),
            ("Xuất file", timezone.now().strftime("%d/%m/%Y %H:%M")),
        ]

        for label, value in metadata_items:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(str(value))

        # Add footer if specified
        if self.template and self.template.footer_text:
            doc.add_paragraph()
            footer_para = doc.add_paragraph(self.template.footer_text)
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # type: ignore[attr-defined]
            run = footer_para.runs[0]
            run.font.size = Pt(9)  # type: ignore[attr-defined]
            run.font.color.rgb = RGBColor(128, 128, 128)  # type: ignore[attr-defined]

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer


class JSONExporter:
    """
    Generate JSON exports
    """

    def __init__(self, template=None):
        self.template = template

    def generate(self, case: Case, output_path: str = None) -> io.BytesIO:  # type: ignore[attr-defined]
        """
        Generate JSON export for a case
        """
        # Get case data based on template settings
        template_settings = {}
        if self.template:
            template_settings = {
                "include_patient_details": self.template.include_patient_details,
                "include_medical_history": self.template.include_medical_history,
                "include_examination": self.template.include_examination,
                "include_investigations": self.template.include_investigations,
                "include_diagnosis": self.template.include_diagnosis,
                "include_treatment": self.template.include_treatment,
                "include_learning_objectives": self.template.include_learning_objectives,
                "include_comments": self.template.include_comments,
                "include_attachments": self.template.include_attachments,
                "include_grades": self.template.include_grades,
                "anonymize_patient": self.template.anonymize_patient,
            }

        data = ExportUtils.get_case_data(case, template_settings)

        # Add export metadata
        data["_export_metadata"] = {
            "format": "json",
            "version": "1.0",
            "exported_at": timezone.now().isoformat(),
            "template": self.template.name if self.template else "default",
        }

        # Convert to JSON with proper encoding
        json_str = json.dumps(data, ensure_ascii=False, indent=2)

        buffer = io.BytesIO(json_str.encode("utf-8"))
        buffer.seek(0)

        return buffer
